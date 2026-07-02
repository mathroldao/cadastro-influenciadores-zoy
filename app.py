
import os
import uuid
from io import BytesIO
from datetime import datetime
from urllib.parse import urlencode

import pandas as pd
import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
from docx import Document


# =========================
# CONFIGURAÇÕES
# =========================

APP_NAME = "Cadastro de Influenciadores | Zoy"
APP_ICON = "📄"

SHEET_NAME = "[ZOY] Cadastro Influenciadores | Base"
WORKSHEET_NAME = "Cadastro"

PUBLIC_BASE_URL = "https://cadastro-zoy.streamlit.app/"

ADMIN_USER_DEFAULT = "juridico"
ADMIN_PASS_DEFAULT = "zoy2026"

COLUMNS = [
    "ID",
    "DATA_CADASTRO",
    "STATUS",
    "RAZAO_SOCIAL",
    "CNPJ",
    "REPRESENTANTE_LEGAL",
    "EMAIL",
    "TELEFONE",
    "INSTAGRAM",
    "TIKTOK",
    "CEP",
    "ENDERECO",
    "NUMERO",
    "COMPLEMENTO",
    "BAIRRO",
    "CIDADE",
    "ESTADO",
    "BANCO",
    "AGENCIA",
    "CONTA",
    "TIPO_CONTA",
    "PIX",
    "CLIENTE",
    "CAMPANHA",
    "VALOR_CACHE",
    "ESCOPO",
    "DIREITO_IMAGEM",
    "EXCLUSIVIDADE",
    "REPOST",
    "IMPULSIONAMENTO",
    "CONTRATO_GERADO",
    "CARTA_GERADA",
    "LINK_CONTRATO",
    "LINK_CARTA",
    "DATA_FORMALIZACAO",
    "D4SIGN_STATUS",
    "D4SIGN_ENVELOPE_ID",
    "OBSERVACOES",
]


# =========================
# PÁGINA / ESTILO
# =========================

st.set_page_config(
    page_title=APP_NAME,
    page_icon=APP_ICON,
    layout="wide",
)

st.markdown(
    """
<style>
    :root {
        --zoy-purple: #7C3AED;
        --zoy-purple-dark: #6D28D9;
        --zoy-bg: #FFFFFF;
        --zoy-soft: #F7F4FF;
        --zoy-border: #E5E7EB;
        --zoy-text: #242434;
        --zoy-muted: #6B7280;
    }

    .block-container {
        max-width: 1180px;
        padding-top: 2.2rem;
        padding-bottom: 4rem;
    }

    section[data-testid="stSidebar"] {
        display: none;
    }

    h1, h2, h3 {
        color: var(--zoy-text);
        letter-spacing: -0.03em;
    }

    h1 {
        font-size: 3.2rem !important;
        font-weight: 800 !important;
        margin-bottom: 0.4rem;
    }

    h2 {
        font-size: 2rem !important;
        font-weight: 800 !important;
        margin-top: 1rem;
    }

    h3 {
        font-size: 1.35rem !important;
        font-weight: 750 !important;
    }

    .zoy-hero {
        background: linear-gradient(135deg, #ffffff 0%, #F4EDFF 100%);
        border: 1px solid #E9D5FF;
        padding: 26px 28px;
        border-radius: 22px;
        margin-bottom: 22px;
    }

    .zoy-subtitle {
        color: var(--zoy-muted);
        font-size: 1.05rem;
        margin-top: 0;
        margin-bottom: 0;
    }

    .zoy-card {
        border: 1px solid var(--zoy-border);
        border-radius: 18px;
        padding: 24px;
        background: white;
        margin-bottom: 20px;
        box-shadow: 0 8px 24px rgba(17,24,39,.035);
    }

    .zoy-card-soft {
        border: 1px solid #E9D5FF;
        border-radius: 18px;
        padding: 22px;
        background: #FAF7FF;
        margin-bottom: 18px;
    }

    .zoy-mini-card {
        border: 1px solid var(--zoy-border);
        border-radius: 14px;
        padding: 16px;
        background: #FFFFFF;
        height: 100%;
    }

    .zoy-label {
        color: #6B7280;
        font-size: 0.82rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        margin-bottom: 2px;
    }

    .zoy-value {
        font-size: 1rem;
        color: #1F2937;
        font-weight: 600;
        margin-bottom: 10px;
        word-break: break-word;
    }

    .zoy-pill {
        display: inline-block;
        padding: 5px 11px;
        border-radius: 999px;
        background: #F4EDFF;
        color: #6D28D9;
        font-size: .85rem;
        font-weight: 700;
        border: 1px solid #E9D5FF;
    }

    .zoy-admin-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 18px;
        margin-bottom: 18px;
    }

    .zoy-login-box {
        max-width: 480px;
        margin: 70px auto 0 auto;
        border: 1px solid var(--zoy-border);
        padding: 34px;
        border-radius: 22px;
        box-shadow: 0 18px 42px rgba(17,24,39,.07);
        background: white;
    }

    div[data-testid="stForm"] {
        border: 1px solid var(--zoy-border);
        border-radius: 18px;
        padding: 24px;
        background: white;
    }

    .stTextInput input,
    .stNumberInput input,
    .stTextArea textarea,
    .stSelectbox div[data-baseweb="select"] {
        border-radius: 12px !important;
        background-color: #F3F4F6 !important;
        border: 1px solid transparent !important;
    }

    .stTextInput input:focus,
    .stTextArea textarea:focus {
        border: 1px solid #A855F7 !important;
        box-shadow: 0 0 0 1px #A855F7 !important;
    }

    .stButton > button,
    .stDownloadButton > button,
    button[kind="primary"] {
        background: var(--zoy-purple) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        min-height: 45px !important;
        font-weight: 750 !important;
        box-shadow: 0 8px 18px rgba(124,58,237,.18);
    }

    .stButton > button:hover,
    .stDownloadButton > button:hover {
        background: var(--zoy-purple-dark) !important;
        color: white !important;
        border: none !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        border-bottom: 1px solid var(--zoy-border);
    }

    .stTabs [data-baseweb="tab"] {
        border-radius: 999px 999px 0 0;
        padding: 10px 16px;
        font-weight: 700;
    }

    hr {
        margin: 1.4rem 0;
    }

    @media (max-width: 768px) {
        h1 {
            font-size: 2.2rem !important;
        }
        .block-container {
            padding-top: 1.2rem;
        }
    }
</style>
""",
    unsafe_allow_html=True,
)


# =========================
# HELPERS
# =========================

def get_admin_credentials():
    user = ADMIN_USER_DEFAULT
    password = ADMIN_PASS_DEFAULT

    try:
        user = st.secrets.get("admin", {}).get("user", ADMIN_USER_DEFAULT)
        password = st.secrets.get("admin", {}).get("password", ADMIN_PASS_DEFAULT)
    except Exception:
        pass

    return user, password


def get_query_param(name, default=""):
    try:
        value = st.query_params.get(name, default)
        if isinstance(value, list):
            return value[0] if value else default
        return value
    except Exception:
        params = st.experimental_get_query_params()
        value = params.get(name, [default])
        return value[0] if value else default


def is_admin_route():
    return get_query_param("admin", "") == "zoy"


def force_public_url(params):
    query = urlencode(params)
    return f"{PUBLIC_BASE_URL}?{query}"


def format_money(value):
    try:
        return f"{float(value):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return str(value or "")


def normalize_money_for_input(value):
    if value in [None, ""]:
        return 0.0
    try:
        clean = str(value).replace("R$", "").strip().replace(".", "").replace(",", ".")
        return float(clean)
    except Exception:
        return 0.0


def safe_value(row, key):
    value = row.get(key, "")
    if pd.isna(value):
        return ""
    return value


def render_field(label, value):
    st.markdown(
        f"""
        <div class="zoy-label">{label}</div>
        <div class="zoy-value">{value if value not in [None, ""] else "—"}</div>
        """,
        unsafe_allow_html=True,
    )


# =========================
# GOOGLE SHEETS
# =========================

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    credentials = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=SCOPES,
    )
    return gspread.authorize(credentials)


def get_worksheet():
    client = get_gspread_client()
    sheet = client.open(SHEET_NAME)

    try:
        ws = sheet.worksheet(WORKSHEET_NAME)
    except Exception:
        ws = sheet.add_worksheet(title=WORKSHEET_NAME, rows=1000, cols=60)

    headers = ws.row_values(1)
    if not headers:
        ws.append_row(COLUMNS)

    return ws


def list_records():
    ws = get_worksheet()
    return ws.get_all_records()


def save_record(data):
    ws = get_worksheet()
    ws.append_row([data.get(column, "") for column in COLUMNS])


# =========================
# DOCUMENTOS
# =========================

def find_template(filename):
    paths = [
        os.path.join("templates", filename),
        filename,
    ]

    for path in paths:
        if os.path.exists(path):
            return path

    return None


def replace_in_docx(template_path, data):
    doc = Document(template_path)
    mapping = {f"{{{{{key}}}}}": str(value) for key, value in data.items()}

    def replace_paragraph(paragraph):
        original_text = paragraph.text
        new_text = original_text

        for placeholder, value in mapping.items():
            new_text = new_text.replace(placeholder, value)

        if new_text != original_text:
            paragraph.clear()
            paragraph.add_run(new_text)

    for paragraph in doc.paragraphs:
        replace_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph(paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================
# LOGIN JURÍDICO
# =========================

def login_page():
    st.markdown(
        """
        <div class="zoy-login-box">
            <h1 style="font-size:2.2rem!important;margin-bottom:6px;">Cadastro de Influenciadores | Zoy</h1>
            <p class="zoy-subtitle">Área interna do time jurídico.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.form("login_form"):
        st.write("### Acesso interno")
        user_input = st.text_input("Usuário")
        pass_input = st.text_input("Senha", type="password")
        submitted = st.form_submit_button("Entrar")

    if submitted:
        admin_user, admin_pass = get_admin_credentials()

        if user_input == admin_user and pass_input == admin_pass:
            st.session_state["admin_logged"] = True
            st.rerun()
        else:
            st.error("Usuário ou senha inválidos.")


# =========================
# PÁGINA PÚBLICA
# =========================

def public_page():
    cliente_url = get_query_param("cliente", "")
    campanha_url = get_query_param("campanha", "")
    valor_url = get_query_param("valor", "0")
    escopo_url = get_query_param("escopo", "")

    direito_url = get_query_param("direito_imagem", "N/A")
    exclusividade_url = get_query_param("exclusividade", "N/A")
    repost_url = get_query_param("repost", "N/A")
    impulsionamento_url = get_query_param("impulsionamento", "N/A")

    st.title("Cadastro de Influenciadores | Zoy")

    st.markdown(
        """
        <div class="zoy-hero">
            <p class="zoy-subtitle">
                Este cadastro será utilizado apenas para coleta de dados cadastrais e formalização interna pela equipe da Zoy.
                O preenchimento leva aproximadamente <b>3 minutos</b>.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.form("form_cadastro"):
        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)
        st.subheader("Dados da Empresa")
        razao = st.text_input("Razão Social *")
        cnpj = st.text_input("CNPJ *")
        representante = st.text_input("Representante Legal *")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)
        st.subheader("Contato")
        email = st.text_input("E-mail *")
        telefone = st.text_input("Telefone / WhatsApp *")
        instagram = st.text_input("Instagram *")
        tiktok = st.text_input("TikTok")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)
        st.subheader("Endereço")
        c1, c2 = st.columns([1, 2])
        with c1:
            cep = st.text_input("CEP")
        with c2:
            endereco = st.text_input("Endereço")
        c3, c4 = st.columns([1, 2])
        with c3:
            numero = st.text_input("Número")
        with c4:
            complemento = st.text_input("Complemento")
        c5, c6, c7 = st.columns([1.3, 1.3, 0.8])
        with c5:
            bairro = st.text_input("Bairro")
        with c6:
            cidade = st.text_input("Cidade")
        with c7:
            estado = st.text_input("Estado")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)
        st.subheader("Dados Bancários")
        banco = st.text_input("Banco")
        c8, c9 = st.columns(2)
        with c8:
            agencia = st.text_input("Agência")
        with c9:
            conta = st.text_input("Conta")
        c10, c11 = st.columns([1, 2])
        with c10:
            tipo_conta = st.selectbox("Tipo de Conta", ["Corrente", "Poupança", "Pagamento"])
        with c11:
            pix = st.text_input("Chave Pix")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)
        st.subheader("Dados da Campanha")
        c12, c13 = st.columns(2)
        with c12:
            cliente = st.text_input("Cliente *", value=cliente_url)
        with c13:
            campanha = st.text_input("Campanha *", value=campanha_url)
        valor = st.number_input(
            "Valor do cachê negociado",
            min_value=0.0,
            step=100.0,
            value=normalize_money_for_input(valor_url),
        )
        escopo = st.text_area("Escopo contratado *", value=escopo_url, height=130)
        c14, c15 = st.columns(2)
        with c14:
            direito = st.text_input("Direito de imagem", value=direito_url)
            repost = st.text_input("Repost", value=repost_url)
        with c15:
            exclusividade = st.text_input("Exclusividade", value=exclusividade_url)
            impulsionamento = st.text_input("Impulsionamento", value=impulsionamento_url)
        st.markdown("</div>", unsafe_allow_html=True)

        confirmacao = st.checkbox("Confirmo que as informações preenchidas estão corretas.")
        submitted = st.form_submit_button("Enviar Cadastro")

    if submitted:
        obrigatorios = {
            "Razão Social": razao,
            "CNPJ": cnpj,
            "Representante Legal": representante,
            "E-mail": email,
            "Telefone": telefone,
            "Instagram": instagram,
            "Cliente": cliente,
            "Campanha": campanha,
            "Escopo": escopo,
        }

        faltando = [field for field, value in obrigatorios.items() if not str(value).strip()]

        if faltando:
            st.error("Preencha os campos obrigatórios: " + ", ".join(faltando))
            return

        if not confirmacao:
            st.error("Confirme que as informações estão corretas antes de enviar.")
            return

        record_id = str(uuid.uuid4())

        data = {
            "ID": record_id,
            "DATA_CADASTRO": datetime.now().strftime("%d/%m/%Y %H:%M"),
            "STATUS": "Pendente",
            "RAZAO_SOCIAL": razao,
            "CNPJ": cnpj,
            "REPRESENTANTE_LEGAL": representante,
            "EMAIL": email,
            "TELEFONE": telefone,
            "INSTAGRAM": instagram,
            "TIKTOK": tiktok,
            "CEP": cep,
            "ENDERECO": endereco,
            "NUMERO": numero,
            "COMPLEMENTO": complemento,
            "BAIRRO": bairro,
            "CIDADE": cidade,
            "ESTADO": estado,
            "BANCO": banco,
            "AGENCIA": agencia,
            "CONTA": conta,
            "TIPO_CONTA": tipo_conta,
            "PIX": pix,
            "CLIENTE": cliente,
            "CAMPANHA": campanha,
            "VALOR_CACHE": format_money(valor),
            "ESCOPO": escopo,
            "DIREITO_IMAGEM": direito,
            "EXCLUSIVIDADE": exclusividade,
            "REPOST": repost,
            "IMPULSIONAMENTO": impulsionamento,
            "CONTRATO_GERADO": "Não",
            "CARTA_GERADA": "Não",
            "LINK_CONTRATO": "",
            "LINK_CARTA": "",
            "DATA_FORMALIZACAO": "",
            "D4SIGN_STATUS": "",
            "D4SIGN_ENVELOPE_ID": "",
            "OBSERVACOES": "",
        }

        save_record(data)

        st.success("Cadastro enviado com sucesso! O time da Zoy recebeu suas informações.")
        st.info(f"Protocolo: {record_id[:8].upper()}")
        st.balloons()


# =========================
# ÁREA JURÍDICA
# =========================

def create_link_page():
    st.subheader("Link para Influenciador")

    with st.form("create_link_form"):
        st.markdown('<div class="zoy-card">', unsafe_allow_html=True)

        c1, c2 = st.columns(2)
        with c1:
            link_cliente = st.text_input("Cliente")
            link_campanha = st.text_input("Campanha")
            link_valor = st.text_input("Valor do cachê")
        with c2:
            link_direito = st.text_input("Direito de imagem", value="N/A")
            link_exclusividade = st.text_input("Exclusividade", value="N/A")
            link_repost = st.text_input("Repost", value="N/A")
            link_impulsionamento = st.text_input("Impulsionamento", value="N/A")

        link_escopo = st.text_area("Escopo contratado", height=130)

        submitted = st.form_submit_button("Gerar Link")
        st.markdown("</div>", unsafe_allow_html=True)

    if submitted:
        params = {
            "cliente": link_cliente,
            "campanha": link_campanha,
            "valor": link_valor,
            "escopo": link_escopo,
            "direito_imagem": link_direito,
            "exclusividade": link_exclusividade,
            "repost": link_repost,
            "impulsionamento": link_impulsionamento,
        }

        link = force_public_url(params)

        st.success("Link criado com sucesso.")
        st.code(link)

        st.markdown(
            f"""
            <a href="{link}" target="_blank">
                <button style="
                    background:#7C3AED;
                    color:white;
                    border:none;
                    border-radius:12px;
                    padding:12px 18px;
                    font-weight:750;
                    cursor:pointer;
                ">Abrir link em nova aba</button>
            </a>
            """,
            unsafe_allow_html=True,
        )


def card_metric(label, value):
    st.markdown(
        f"""
        <div class="zoy-mini-card">
            <div class="zoy-label">{label}</div>
            <div style="font-size:1.9rem;font-weight:800;color:#2B2B36;">{value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def cadastros_page():
    st.subheader("Cadastros")

    try:
        records = list_records()
    except Exception as error:
        st.error("Não foi possível carregar os cadastros.")
        st.exception(error)
        return

    if not records:
        st.info("Nenhum cadastro encontrado.")
        return

    df = pd.DataFrame(records)

    total = len(df)
    pendentes = len(df[df["STATUS"].astype(str).str.lower() == "pendente"]) if "STATUS" in df else 0
    contratos = len(df[df["CONTRATO_GERADO"].astype(str).str.lower() == "sim"]) if "CONTRATO_GERADO" in df else 0
    cartas = len(df[df["CARTA_GERADA"].astype(str).str.lower() == "sim"]) if "CARTA_GERADA" in df else 0

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        card_metric("Total de cadastros", total)
    with m2:
        card_metric("Pendentes", pendentes)
    with m3:
        card_metric("Contratos gerados", contratos)
    with m4:
        card_metric("Cartas geradas", cartas)

    st.write("")

    c1, c2 = st.columns([3, 1])
    with c1:
        search = st.text_input("Buscar por razão social, Instagram, cliente ou campanha")
    with c2:
        status_filter = st.selectbox("Status", ["Todos", "Pendente", "Formalizado", "Finalizado"])

    filtered = df.copy()

    if search:
        text = search.lower()
        filtered = filtered[
            filtered.astype(str)
            .apply(lambda col: col.str.lower().str.contains(text, na=False))
            .any(axis=1)
        ]

    if status_filter != "Todos" and "STATUS" in filtered:
        filtered = filtered[filtered["STATUS"].astype(str) == status_filter]

    if filtered.empty:
        st.info("Nenhum cadastro encontrado com esses filtros.")
        return

    for index, row in filtered.iterrows():
        title = f"{safe_value(row, 'RAZAO_SOCIAL')} | {safe_value(row, 'INSTAGRAM')} | {safe_value(row, 'CLIENTE')} - {safe_value(row, 'CAMPANHA')}"

        with st.expander(title):
            left, right = st.columns(2)

            with left:
                st.markdown('<div class="zoy-card-soft">', unsafe_allow_html=True)
                st.write("### Empresa")
                render_field("Razão Social", safe_value(row, "RAZAO_SOCIAL"))
                render_field("CNPJ", safe_value(row, "CNPJ"))
                render_field("Representante Legal", safe_value(row, "REPRESENTANTE_LEGAL"))
                render_field("E-mail", safe_value(row, "EMAIL"))
                render_field("Telefone", safe_value(row, "TELEFONE"))
                render_field("Instagram", safe_value(row, "INSTAGRAM"))
                render_field("TikTok", safe_value(row, "TIKTOK"))
                st.markdown("</div>", unsafe_allow_html=True)

                st.markdown('<div class="zoy-card-soft">', unsafe_allow_html=True)
                st.write("### Bancário")
                render_field("Banco", safe_value(row, "BANCO"))
                render_field("Agência", safe_value(row, "AGENCIA"))
                render_field("Conta", safe_value(row, "CONTA"))
                render_field("Tipo de Conta", safe_value(row, "TIPO_CONTA"))
                render_field("Pix", safe_value(row, "PIX"))
                st.markdown("</div>", unsafe_allow_html=True)

            with right:
                st.markdown('<div class="zoy-card-soft">', unsafe_allow_html=True)
                st.write("### Campanha")
                render_field("Cliente", safe_value(row, "CLIENTE"))
                render_field("Campanha", safe_value(row, "CAMPANHA"))
                render_field("Valor", f"R$ {safe_value(row, 'VALOR_CACHE')}")
                render_field("Escopo", safe_value(row, "ESCOPO"))
                render_field("Direito de imagem", safe_value(row, "DIREITO_IMAGEM"))
                render_field("Exclusividade", safe_value(row, "EXCLUSIVIDADE"))
                render_field("Repost", safe_value(row, "REPOST"))
                render_field("Impulsionamento", safe_value(row, "IMPULSIONAMENTO"))
                st.markdown("</div>", unsafe_allow_html=True)

                st.markdown('<div class="zoy-card-soft">', unsafe_allow_html=True)
                st.write("### Endereço")
                render_field("CEP", safe_value(row, "CEP"))
                render_field("Endereço", f"{safe_value(row, 'ENDERECO')}, {safe_value(row, 'NUMERO')} {safe_value(row, 'COMPLEMENTO')}")
                render_field("Bairro", safe_value(row, "BAIRRO"))
                render_field("Cidade/Estado", f"{safe_value(row, 'CIDADE')} / {safe_value(row, 'ESTADO')}")
                st.markdown("</div>", unsafe_allow_html=True)

            st.write("")
            st.write("### Documentos")

            document_data = row.to_dict()
            document_data["DATA_EXTENSO"] = datetime.now().strftime("%d/%m/%Y")

            contrato_path = find_template("contrato.docx")
            carta_path = find_template("carta.docx")

            d1, d2 = st.columns(2)

            with d1:
                if contrato_path:
                    contrato_buffer = replace_in_docx(contrato_path, document_data)
                    st.download_button(
                        "Baixar Contrato",
                        contrato_buffer,
                        file_name=f"Contrato - {safe_value(row, 'RAZAO_SOCIAL')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"contrato_{safe_value(row, 'ID')}_{index}",
                    )
                else:
                    st.warning("Arquivo contrato.docx não encontrado.")

            with d2:
                if carta_path:
                    carta_buffer = replace_in_docx(carta_path, document_data)
                    st.download_button(
                        "Baixar Carta",
                        carta_buffer,
                        file_name=f"Carta - {safe_value(row, 'RAZAO_SOCIAL')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"carta_{safe_value(row, 'ID')}_{index}",
                    )
                else:
                    st.warning("Arquivo carta.docx não encontrado.")


def juridico_page():
    if not st.session_state.get("admin_logged"):
        login_page()
        return

    st.markdown(
        """
        <div class="zoy-admin-header">
            <div>
                <h1 style="margin-bottom:0;">Cadastro de Influenciadores</h1>
                <p class="zoy-subtitle">Área interna da Zoy para gestão de cadastros e documentos.</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    top1, top2, top3 = st.columns([1.2, 1.2, 5])

    if "admin_view" not in st.session_state:
        st.session_state["admin_view"] = "cadastros"

    with top1:
        if st.button("Cadastros", use_container_width=True):
            st.session_state["admin_view"] = "cadastros"
            st.rerun()

    with top2:
        if st.button("Link para Influenciador", use_container_width=True):
            st.session_state["admin_view"] = "link"
            st.rerun()

    with top3:
        logout = st.button("Sair", use_container_width=False)
        if logout:
            st.session_state["admin_logged"] = False
            st.rerun()

    st.divider()

    if st.session_state["admin_view"] == "link":
        create_link_page()
    else:
        cadastros_page()


# =========================
# ROUTER
# =========================

if is_admin_route():
    juridico_page()
else:
    public_page()
